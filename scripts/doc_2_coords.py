# pip install pdf2docx -q
# pip install python-docx -q

import os
import sys
import docx
import argparse

import numpy as np
import pandas as pd
import matplotlib.pyplot as plt

from glob import glob
from pdf2docx import Converter

def read_docx_with_tables(filepath):
    """
    Reads text content from a .docx file, including paragraphs and tables.

    Args:
        filepath (str): The path to the .docx file.

    Returns:
        str: The extracted text content from paragraphs and tables,
             with elements joined by newlines. Table content is formatted
             to show rows and cells separated by tabs.
             Returns None if the file cannot be read or doesn't exist.
    """
    if not os.path.exists(filepath):
        print(f"Error: File not found at {filepath}")
        return None
    if not filepath.lower().endswith('.docx'):
        print(f"Error: File {filepath} is not a .docx file.")
        return None

    try:
        doc = docx.Document(filepath)
        full_content = []
        base = []
        coords = []

        utmx = 0
        utmy = 0

        # Iterate through paragraphs and tables in document order (approximately)
        # The Document object stores paragraphs and tables separately.
        # A more complex approach involves parsing the underlying XML (_body._element.xpath('.//w:p | .//w:tbl'))
        # to get perfect order, but iterating through them separately is often sufficient.

        # print("Reading paragraphs...")
        for para in doc.paragraphs:
            if para.text.strip(): # Add non-empty paragraphs
                 # full_content.append(para.text)
                 pass

        print(f"Found {len(doc.tables)} table(s). Reading tables...")
        for i, table in enumerate(doc.tables):
            # print(f"  Reading Table {i+1}...")
            # full_content.append(f"--- Table {i+1} Start ---")
            for row in table.rows:
                row_cells = []
                for cell in row.cells:
                    # Clean up cell text (replace newlines within a cell with spaces)
                    # print(cell.text)
                    cell_text = cell.text.replace('\n', ' ').strip()
                    # print(cell_text)
                    row_cells.append(cell_text)
                    
                    if "X" in cell_text:utmx+=1
                    elif "Y" in cell_text:utmy+=1

                # Join cells of a row with a Tab character for basic structure
                # row_text = "\t|\t".join(row_cells)
                row_text = ",".join(row_cells)
                # print(row_text)
                if (
                    ("ipu" not in row_text) 
                    and ("X" not in row_text)
                    and ("Y" not in row_text)
                ):
                    is_valid = True
                    fields = row_text.split(',')

                    for field in fields:
                        # Remove leading/trailing whitespace for robustness
                        trimmed_field = field.strip()
                        # Check if the field is NOT empty AND contains anything other than digits
                        if trimmed_field != '' and not trimmed_field.isdecimal():
                            is_valid = False
                            break # Found an invalid field, no need to check the rest of the line
                    
                    if is_valid:
                        full_content.append(row_text)
                        coords.append(row_cells)
                        ### print(row_cells)
                        # print(row_text)
                        
                # print()
                # full_content.append("\t|\t".join(row_cells))
            # full_content.append(f"--- Table {i+1} End ---")
        
        if utmx==0 and utmy==0:
            print("❌ No se encontraron coordenadas en el documento.")
            return "", []
        else:
            print("✅ Se encontraron coordenadas en el documento.")

        if len(coords)==0:
            for i, table in enumerate(doc.tables):
                print(f"  Reading Table {i+1}...")
                # full_content.append(f"--- Table {i+1} Start ---")
                for row in table.rows:
                    for cell in row.cells:
                        cell_text0 = cell.text.replace(' ', '').strip()
                        cell_text = cell.text.strip()
                        is_num = cell_text0.isnumeric()
                        # print(cell.text)
                        if (is_num 
                            #and cell_text.count(" ")>=2
                            ):
                            # print(cell.text)
                            base.append(cell.text)
        
            idx_2 = base.index("2")
            # display(base)
            # print()
            base2 = [i for i in base[:idx_2] if len(i.split(" ")[0])>=6]
            # display(base2)
            n_b = len(base2)

            cont = 1
            for i in range(0, n_b, 2):
                c_c = base2[i]
                c_d = base2[i+1]

                datax = c_c.split(" ")
                datay = c_d.split(" ")

                for i, dx in enumerate(datax):
                    coords.append([cont,dx,datay[i]])
                    full_content.append(f"{cont},{dx},{datay[i]}")
                    cont+=1

            # print(n_b)
            # print(coords)
            if len(coords)>0:
                return '\n'.join(full_content), coords
            else:
                print("😳 Please, take a look down here.")
                return "", None
            # return full_content, coords

        else:
            return '\n'.join(full_content), coords

    # Catches errors like file not found (if os check fails somehow)
    # or if the file is not a valid docx format (e.g., corrupted, password-protected, or older .doc)
    except Exception as e:

        exc_type, exc_obj, exc_tb = sys.exc_info()
        error_l = exc_tb.tb_lineno
        cadena_error = str(exc_type) + " => " + str(exc_obj)

        # print(f"Error reading {filepath}: {e}")
        print(f"Error line: {error_l}")
        print(f"Error message: {cadena_error}")
        return None, None
